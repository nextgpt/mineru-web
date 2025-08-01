"""
文档导出服务
支持PDF、Word、Markdown等格式导出，添加导出模板和样式配置
"""
from typing import Optional, Dict, Any, List
from sqlalchemy.orm import Session
from app.models.project import Project
from app.models.bid_document import BidDocument
from app.models.bid_outline import BidOutline
from app.models.requirement_analysis import RequirementAnalysis
import logging
import io
import os
import tempfile
from datetime import datetime
from enum import Enum

logger = logging.getLogger(__name__)


class ExportFormat(Enum):
    """导出格式枚举"""
    PDF = "pdf"
    WORD = "docx"
    MARKDOWN = "md"
    HTML = "html"
    TXT = "txt"


class DocumentExportService:
    """文档导出服务"""
    
    def __init__(self):
        """初始化文档导出服务"""
        self.supported_formats = [ExportFormat.PDF, ExportFormat.WORD, ExportFormat.MARKDOWN, ExportFormat.HTML, ExportFormat.TXT]
        self.export_templates = self._load_export_templates()
    
    def export_project_document(self, db: Session, project_id: int, user_id: str, 
                               export_format: ExportFormat, include_outline: bool = True,
                               include_analysis: bool = True) -> Optional[bytes]:
        """
        导出项目完整文档
        
        Args:
            db: 数据库会话
            project_id: 项目ID
            user_id: 用户ID
            export_format: 导出格式
            include_outline: 是否包含大纲
            include_analysis: 是否包含需求分析
            
        Returns:
            导出的文档字节数据
        """
        try:
            logger.info(f"开始导出项目 {project_id} 的文档，格式: {export_format.value}")
            
            # 获取项目信息
            project = db.query(Project).filter(
                Project.id == project_id,
                Project.user_id == user_id
            ).first()
            
            if not project:
                logger.error(f"项目 {project_id} 不存在")
                return None
            
            # 获取文档内容
            content = self._build_export_content(db, project, include_outline, include_analysis)
            
            if not content:
                logger.error(f"项目 {project_id} 没有可导出的内容")
                return None
            
            # 根据格式导出
            if export_format == ExportFormat.MARKDOWN:
                return self._export_to_markdown(content, project)
            elif export_format == ExportFormat.HTML:
                return self._export_to_html(content, project)
            elif export_format == ExportFormat.TXT:
                return self._export_to_txt(content, project)
            elif export_format == ExportFormat.PDF:
                return self._export_to_pdf(content, project)
            elif export_format == ExportFormat.WORD:
                return self._export_to_word(content, project)
            else:
                logger.error(f"不支持的导出格式: {export_format.value}")
                return None
                
        except Exception as e:
            logger.error(f"导出文档失败: {str(e)}")
            return None
    
    def export_single_document(self, db: Session, document_id: int, user_id: str, 
                              export_format: ExportFormat) -> Optional[bytes]:
        """
        导出单个文档
        
        Args:
            db: 数据库会话
            document_id: 文档ID
            user_id: 用户ID
            export_format: 导出格式
            
        Returns:
            导出的文档字节数据
        """
        try:
            # 获取文档
            document = db.query(BidDocument).filter(
                BidDocument.id == document_id,
                BidDocument.user_id == user_id
            ).first()
            
            if not document:
                logger.error(f"文档 {document_id} 不存在")
                return None
            
            # 构建单文档内容
            content = {
                'title': document.title,
                'content': document.content,
                'created_at': document.created_at,
                'version': document.version
            }
            
            # 获取项目信息
            project = db.query(Project).filter(Project.id == document.project_id).first()
            
            # 根据格式导出
            if export_format == ExportFormat.MARKDOWN:
                return self._export_single_to_markdown(content, project)
            elif export_format == ExportFormat.HTML:
                return self._export_single_to_html(content, project)
            elif export_format == ExportFormat.TXT:
                return self._export_single_to_txt(content, project)
            elif export_format == ExportFormat.PDF:
                return self._export_single_to_pdf(content, project)
            elif export_format == ExportFormat.WORD:
                return self._export_single_to_word(content, project)
            else:
                logger.error(f"不支持的导出格式: {export_format.value}")
                return None
                
        except Exception as e:
            logger.error(f"导出单个文档失败: {str(e)}")
            return None
    
    def _build_export_content(self, db: Session, project: Project, 
                            include_outline: bool, include_analysis: bool) -> Optional[Dict[str, Any]]:
        """构建导出内容"""
        try:
            content = {
                'project': {
                    'name': project.name,
                    'description': project.description or '',
                    'status': project.status.value,
                    'created_at': project.created_at,
                    'updated_at': project.updated_at
                },
                'sections': []
            }
            
            # 添加需求分析
            if include_analysis:
                analysis = db.query(RequirementAnalysis).filter(
                    RequirementAnalysis.project_id == project.id,
                    RequirementAnalysis.user_id == project.user_id
                ).first()
                
                if analysis:
                    content['analysis'] = {
                        'project_overview': analysis.project_overview or '',
                        'client_info': analysis.client_info or '',
                        'budget_info': analysis.budget_info or '',
                        'critical_requirements': analysis.critical_requirements or '',
                        'important_requirements': analysis.important_requirements or '',
                        'general_requirements': analysis.general_requirements or ''
                    }
            
            # 获取大纲和文档
            outlines = db.query(BidOutline).filter(
                BidOutline.project_id == project.id,
                BidOutline.user_id == project.user_id
            ).order_by(BidOutline.level, BidOutline.order_index).all()
            
            documents = db.query(BidDocument).filter(
                BidDocument.project_id == project.id,
                BidDocument.user_id == project.user_id
            ).all()
            
            # 创建大纲ID到文档的映射
            doc_map = {doc.outline_id: doc for doc in documents if doc.outline_id}
            
            # 构建章节内容
            for outline in outlines:
                section = {
                    'sequence': outline.sequence,
                    'title': outline.title,
                    'level': outline.level,
                    'outline_content': outline.content or '',
                    'document_content': '',
                    'has_document': False
                }
                
                if outline.id in doc_map:
                    document = doc_map[outline.id]
                    section['document_content'] = document.content
                    section['has_document'] = True
                    section['document_status'] = document.status.value
                    section['document_version'] = document.version
                
                content['sections'].append(section)
            
            return content
            
        except Exception as e:
            logger.error(f"构建导出内容失败: {str(e)}")
            return None
    
    def _export_to_markdown(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出为Markdown格式"""
        try:
            md_lines = []
            
            # 文档标题
            md_lines.append(f"# {content['project']['name']}")
            md_lines.append("")
            
            # 项目信息
            if content['project']['description']:
                md_lines.append(f"**项目描述：** {content['project']['description']}")
                md_lines.append("")
            
            md_lines.append(f"**项目状态：** {content['project']['status']}")
            md_lines.append(f"**创建时间：** {content['project']['created_at'].strftime('%Y-%m-%d %H:%M:%S')}")
            md_lines.append(f"**导出时间：** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            md_lines.append("")
            md_lines.append("---")
            md_lines.append("")
            
            # 需求分析
            if 'analysis' in content:
                analysis = content['analysis']
                md_lines.append("## 需求分析")
                md_lines.append("")
                
                if analysis['project_overview']:
                    md_lines.append("### 项目概览")
                    md_lines.append("")
                    md_lines.append(analysis['project_overview'])
                    md_lines.append("")
                
                if analysis['client_info']:
                    md_lines.append("### 甲方信息")
                    md_lines.append("")
                    md_lines.append(analysis['client_info'])
                    md_lines.append("")
                
                if analysis['budget_info']:
                    md_lines.append("### 预算信息")
                    md_lines.append("")
                    md_lines.append(analysis['budget_info'])
                    md_lines.append("")
                
                if analysis['critical_requirements']:
                    md_lines.append("### 关键需求")
                    md_lines.append("")
                    md_lines.append(analysis['critical_requirements'])
                    md_lines.append("")
                
                if analysis['important_requirements']:
                    md_lines.append("### 重要需求")
                    md_lines.append("")
                    md_lines.append(analysis['important_requirements'])
                    md_lines.append("")
                
                if analysis['general_requirements']:
                    md_lines.append("### 一般需求")
                    md_lines.append("")
                    md_lines.append(analysis['general_requirements'])
                    md_lines.append("")
                
                md_lines.append("---")
                md_lines.append("")
            
            # 目录
            md_lines.append("## 目录")
            md_lines.append("")
            for section in content['sections']:
                indent = "  " * (section['level'] - 1)
                md_lines.append(f"{indent}- [{section['sequence']} {section['title']}](#{section['sequence'].replace('.', '')}-{section['title'].replace(' ', '-').lower()})")
            md_lines.append("")
            md_lines.append("---")
            md_lines.append("")
            
            # 章节内容
            for section in content['sections']:
                level_prefix = "#" * (section['level'] + 1)
                anchor = f"{section['sequence'].replace('.', '')}-{section['title'].replace(' ', '-').lower()}"
                md_lines.append(f"{level_prefix} {section['sequence']} {section['title']} {{#{anchor}}}")
                md_lines.append("")
                
                if section['has_document'] and section['document_content']:
                    md_lines.append(section['document_content'])
                else:
                    md_lines.append("*此章节内容尚未生成*")
                
                md_lines.append("")
                md_lines.append("---")
                md_lines.append("")
            
            return "\n".join(md_lines).encode('utf-8')
            
        except Exception as e:
            logger.error(f"导出Markdown失败: {str(e)}")
            return b""
    
    def _export_to_html(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出为HTML格式"""
        try:
            html_template = self.export_templates.get('html', {}).get('template', '')
            
            if not html_template:
                # 使用默认HTML模板
                html_template = self._get_default_html_template()
            
            # 构建HTML内容
            html_content = self._build_html_content(content)
            
            # 替换模板变量
            html_output = html_template.format(
                title=content['project']['name'],
                content=html_content,
                export_time=datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            )
            
            return html_output.encode('utf-8')
            
        except Exception as e:
            logger.error(f"导出HTML失败: {str(e)}")
            return b""
    
    def _export_to_txt(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出为纯文本格式"""
        try:
            txt_lines = []
            
            # 文档标题
            txt_lines.append("=" * 60)
            txt_lines.append(f"标书文档：{content['project']['name']}")
            txt_lines.append("=" * 60)
            txt_lines.append("")
            
            # 项目信息
            txt_lines.append("项目信息：")
            txt_lines.append("-" * 30)
            if content['project']['description']:
                txt_lines.append(f"项目描述：{content['project']['description']}")
            txt_lines.append(f"项目状态：{content['project']['status']}")
            txt_lines.append(f"创建时间：{content['project']['created_at'].strftime('%Y-%m-%d %H:%M:%S')}")
            txt_lines.append(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            txt_lines.append("")
            
            # 需求分析
            if 'analysis' in content:
                analysis = content['analysis']
                txt_lines.append("需求分析：")
                txt_lines.append("-" * 30)
                
                if analysis['project_overview']:
                    txt_lines.append("项目概览：")
                    txt_lines.append(analysis['project_overview'])
                    txt_lines.append("")
                
                if analysis['client_info']:
                    txt_lines.append("甲方信息：")
                    txt_lines.append(analysis['client_info'])
                    txt_lines.append("")
                
                if analysis['critical_requirements']:
                    txt_lines.append("关键需求：")
                    txt_lines.append(analysis['critical_requirements'])
                    txt_lines.append("")
            
            # 章节内容
            txt_lines.append("标书内容：")
            txt_lines.append("-" * 30)
            
            for section in content['sections']:
                # 章节标题
                title_line = f"{section['sequence']} {section['title']}"
                txt_lines.append(title_line)
                txt_lines.append("-" * len(title_line))
                
                if section['has_document'] and section['document_content']:
                    txt_lines.append(section['document_content'])
                else:
                    txt_lines.append("此章节内容尚未生成")
                
                txt_lines.append("")
                txt_lines.append("")
            
            return "\n".join(txt_lines).encode('utf-8')
            
        except Exception as e:
            logger.error(f"导出TXT失败: {str(e)}")
            return b""
    
    def _export_to_pdf(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出为PDF格式"""
        try:
            # 首先生成HTML内容
            html_content = self._export_to_html(content, project)
            
            # 使用weasyprint将HTML转换为PDF
            try:
                import weasyprint
                pdf_bytes = weasyprint.HTML(string=html_content.decode('utf-8')).write_pdf()
                return pdf_bytes
            except ImportError:
                logger.warning("weasyprint未安装，尝试使用替代方案")
                # 如果weasyprint不可用，返回HTML内容作为备选
                return html_content
            
        except Exception as e:
            logger.error(f"导出PDF失败: {str(e)}")
            return b""
    
    def _export_to_word(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出为Word格式"""
        try:
            # 使用python-docx生成Word文档
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # 设置文档标题
                title = doc.add_heading(content['project']['name'], 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加项目信息
                doc.add_heading('项目信息', level=1)
                if content['project']['description']:
                    p = doc.add_paragraph()
                    p.add_run('项目描述：').bold = True
                    p.add_run(content['project']['description'])
                
                p = doc.add_paragraph()
                p.add_run('项目状态：').bold = True
                p.add_run(content['project']['status'])
                
                p = doc.add_paragraph()
                p.add_run('创建时间：').bold = True
                p.add_run(content['project']['created_at'].strftime('%Y-%m-%d %H:%M:%S'))
                
                p = doc.add_paragraph()
                p.add_run('导出时间：').bold = True
                p.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
                
                # 添加需求分析
                if 'analysis' in content:
                    analysis = content['analysis']
                    doc.add_heading('需求分析', level=1)
                    
                    if analysis['project_overview']:
                        doc.add_heading('项目概览', level=2)
                        doc.add_paragraph(analysis['project_overview'])
                    
                    if analysis['client_info']:
                        doc.add_heading('甲方信息', level=2)
                        doc.add_paragraph(analysis['client_info'])
                    
                    if analysis['critical_requirements']:
                        doc.add_heading('关键需求', level=2)
                        doc.add_paragraph(analysis['critical_requirements'])
                
                # 添加章节内容
                for section in content['sections']:
                    doc.add_heading(f"{section['sequence']} {section['title']}", level=section['level'])
                    
                    if section['has_document'] and section['document_content']:
                        doc.add_paragraph(section['document_content'])
                    else:
                        p = doc.add_paragraph('此章节内容尚未生成')
                        p.italic = True
                
                # 保存到内存
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                return doc_io.getvalue()
                
            except ImportError:
                logger.warning("python-docx未安装，无法生成Word文档")
                # 返回纯文本作为备选
                return self._export_to_txt(content, project)
            
        except Exception as e:
            logger.error(f"导出Word失败: {str(e)}")
            return b""
    
    def _export_single_to_markdown(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出单个文档为Markdown"""
        try:
            md_lines = []
            
            md_lines.append(f"# {content['title']}")
            md_lines.append("")
            
            if project:
                md_lines.append(f"**所属项目：** {project.name}")
                md_lines.append("")
            
            md_lines.append(f"**文档版本：** v{content['version']}")
            md_lines.append(f"**创建时间：** {content['created_at'].strftime('%Y-%m-%d %H:%M:%S')}")
            md_lines.append(f"**导出时间：** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            md_lines.append("")
            md_lines.append("---")
            md_lines.append("")
            
            md_lines.append("## 内容")
            md_lines.append("")
            md_lines.append(content['content'])
            
            return "\n".join(md_lines).encode('utf-8')
            
        except Exception as e:
            logger.error(f"导出单个文档Markdown失败: {str(e)}")
            return b""
    
    def _export_single_to_html(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出单个文档为HTML"""
        try:
            html_template = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; line-height: 1.6; margin: 40px; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        .meta {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; margin: 20px 0; }}
        .content {{ margin-top: 30px; }}
        pre {{ background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <div class="meta">
        <p><strong>所属项目：</strong> {project_name}</p>
        <p><strong>文档版本：</strong> v{version}</p>
        <p><strong>创建时间：</strong> {created_at}</p>
        <p><strong>导出时间：</strong> {export_time}</p>
    </div>
    <div class="content">
        <h2>内容</h2>
        <div>{content}</div>
    </div>
</body>
</html>
"""
            
            html_output = html_template.format(
                title=content['title'],
                project_name=project.name if project else '未知项目',
                version=content['version'],
                created_at=content['created_at'].strftime('%Y-%m-%d %H:%M:%S'),
                export_time=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                content=content['content'].replace('\n', '<br>')
            )
            
            return html_output.encode('utf-8')
            
        except Exception as e:
            logger.error(f"导出单个文档HTML失败: {str(e)}")
            return b""
    
    def _export_single_to_txt(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出单个文档为TXT"""
        try:
            txt_lines = []
            
            txt_lines.append("=" * 60)
            txt_lines.append(f"文档：{content['title']}")
            txt_lines.append("=" * 60)
            txt_lines.append("")
            
            if project:
                txt_lines.append(f"所属项目：{project.name}")
            txt_lines.append(f"文档版本：v{content['version']}")
            txt_lines.append(f"创建时间：{content['created_at'].strftime('%Y-%m-%d %H:%M:%S')}")
            txt_lines.append(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            txt_lines.append("")
            txt_lines.append("-" * 60)
            txt_lines.append("")
            txt_lines.append("内容：")
            txt_lines.append("")
            txt_lines.append(content['content'])
            
            return "\n".join(txt_lines).encode('utf-8')
            
        except Exception as e:
            logger.error(f"导出单个文档TXT失败: {str(e)}")
            return b""
    
    def _export_single_to_pdf(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出单个文档为PDF"""
        try:
            html_content = self._export_single_to_html(content, project)
            
            try:
                import weasyprint
                pdf_bytes = weasyprint.HTML(string=html_content.decode('utf-8')).write_pdf()
                return pdf_bytes
            except ImportError:
                logger.warning("weasyprint未安装，返回HTML内容")
                return html_content
            
        except Exception as e:
            logger.error(f"导出单个文档PDF失败: {str(e)}")
            return b""
    
    def _export_single_to_word(self, content: Dict[str, Any], project: Project) -> bytes:
        """导出单个文档为Word"""
        try:
            try:
                from docx import Document
                
                doc = Document()
                
                # 文档标题
                doc.add_heading(content['title'], 0)
                
                # 元信息
                if project:
                    p = doc.add_paragraph()
                    p.add_run('所属项目：').bold = True
                    p.add_run(project.name)
                
                p = doc.add_paragraph()
                p.add_run('文档版本：').bold = True
                p.add_run(f"v{content['version']}")
                
                p = doc.add_paragraph()
                p.add_run('创建时间：').bold = True
                p.add_run(content['created_at'].strftime('%Y-%m-%d %H:%M:%S'))
                
                p = doc.add_paragraph()
                p.add_run('导出时间：').bold = True
                p.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
                
                # 内容
                doc.add_heading('内容', level=1)
                doc.add_paragraph(content['content'])
                
                # 保存到内存
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                return doc_io.getvalue()
                
            except ImportError:
                logger.warning("python-docx未安装，返回TXT内容")
                return self._export_single_to_txt(content, project)
            
        except Exception as e:
            logger.error(f"导出单个文档Word失败: {str(e)}")
            return b""
    
    def _build_html_content(self, content: Dict[str, Any]) -> str:
        """构建HTML内容"""
        try:
            html_parts = []
            
            # 项目信息
            html_parts.append('<div class="project-info">')
            if content['project']['description']:
                html_parts.append(f'<p><strong>项目描述：</strong>{content["project"]["description"]}</p>')
            html_parts.append(f'<p><strong>项目状态：</strong>{content["project"]["status"]}</p>')
            html_parts.append(f'<p><strong>创建时间：</strong>{content["project"]["created_at"].strftime("%Y-%m-%d %H:%M:%S")}</p>')
            html_parts.append('</div>')
            
            # 需求分析
            if 'analysis' in content:
                analysis = content['analysis']
                html_parts.append('<div class="analysis-section">')
                html_parts.append('<h2>需求分析</h2>')
                
                if analysis['project_overview']:
                    html_parts.append('<h3>项目概览</h3>')
                    html_parts.append(f'<div class="content-block">{analysis["project_overview"].replace(chr(10), "<br>")}</div>')
                
                if analysis['client_info']:
                    html_parts.append('<h3>甲方信息</h3>')
                    html_parts.append(f'<div class="content-block">{analysis["client_info"].replace(chr(10), "<br>")}</div>')
                
                if analysis['critical_requirements']:
                    html_parts.append('<h3>关键需求</h3>')
                    html_parts.append(f'<div class="content-block">{analysis["critical_requirements"].replace(chr(10), "<br>")}</div>')
                
                html_parts.append('</div>')
            
            # 目录
            html_parts.append('<div class="toc-section">')
            html_parts.append('<h2>目录</h2>')
            html_parts.append('<ul class="toc-list">')
            for section in content['sections']:
                indent_class = f"toc-level-{section['level']}"
                html_parts.append(f'<li class="{indent_class}"><a href="#section-{section["sequence"].replace(".", "-")}">{section["sequence"]} {section["title"]}</a></li>')
            html_parts.append('</ul>')
            html_parts.append('</div>')
            
            # 章节内容
            html_parts.append('<div class="sections-content">')
            for section in content['sections']:
                section_id = f"section-{section['sequence'].replace('.', '-')}"
                level_class = f"section-level-{section['level']}"
                
                html_parts.append(f'<div class="section {level_class}" id="{section_id}">')
                html_parts.append(f'<h{section["level"] + 1}>{section["sequence"]} {section["title"]}</h{section["level"] + 1}>')
                
                if section['has_document'] and section['document_content']:
                    html_parts.append(f'<div class="section-content">{section["document_content"].replace(chr(10), "<br>")}</div>')
                else:
                    html_parts.append('<div class="section-content placeholder">此章节内容尚未生成</div>')
                
                html_parts.append('</div>')
            html_parts.append('</div>')
            
            return '\n'.join(html_parts)
            
        except Exception as e:
            logger.error(f"构建HTML内容失败: {str(e)}")
            return ""
    
    def _get_default_html_template(self) -> str:
        """获取默认HTML模板"""
        return """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Microsoft YaHei', 'SimSun', Arial, sans-serif;
            line-height: 1.8;
            margin: 0;
            padding: 40px;
            background-color: #ffffff;
            color: #333333;
        }}
        
        h1 {{
            color: #2c3e50;
            text-align: center;
            border-bottom: 3px solid #3498db;
            padding-bottom: 20px;
            margin-bottom: 40px;
            font-size: 2.5em;
        }}
        
        h2 {{
            color: #34495e;
            border-left: 4px solid #3498db;
            padding-left: 15px;
            margin-top: 40px;
            margin-bottom: 20px;
            font-size: 1.8em;
        }}
        
        h3 {{
            color: #2c3e50;
            margin-top: 30px;
            margin-bottom: 15px;
            font-size: 1.4em;
        }}
        
        .project-info {{
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 8px;
            margin: 30px 0;
            border-left: 4px solid #17a2b8;
        }}
        
        .analysis-section {{
            margin: 40px 0;
        }}
        
        .content-block {{
            background-color: #ffffff;
            padding: 20px;
            border: 1px solid #e9ecef;
            border-radius: 5px;
            margin: 15px 0;
            white-space: pre-wrap;
        }}
        
        .toc-section {{
            background-color: #f1f3f4;
            padding: 25px;
            border-radius: 8px;
            margin: 30px 0;
        }}
        
        .toc-list {{
            list-style: none;
            padding-left: 0;
        }}
        
        .toc-list li {{
            margin: 8px 0;
            padding: 5px 0;
        }}
        
        .toc-level-1 {{
            font-weight: bold;
            font-size: 1.1em;
        }}
        
        .toc-level-2 {{
            padding-left: 20px;
            font-size: 1.0em;
        }}
        
        .toc-level-3 {{
            padding-left: 40px;
            font-size: 0.9em;
            color: #666;
        }}
        
        .toc-list a {{
            text-decoration: none;
            color: #2c3e50;
            transition: color 0.3s;
        }}
        
        .toc-list a:hover {{
            color: #3498db;
        }}
        
        .sections-content {{
            margin-top: 40px;
        }}
        
        .section {{
            margin: 40px 0;
            padding: 20px 0;
            border-bottom: 1px solid #e9ecef;
        }}
        
        .section-content {{
            margin: 20px 0;
            padding: 20px;
            background-color: #ffffff;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            white-space: pre-wrap;
        }}
        
        .placeholder {{
            font-style: italic;
            color: #6c757d;
            background-color: #f8f9fa;
        }}
        
        .export-info {{
            text-align: center;
            margin-top: 50px;
            padding: 20px;
            background-color: #e9ecef;
            border-radius: 5px;
            font-size: 0.9em;
            color: #6c757d;
        }}
        
        @media print {{
            body {{ margin: 20px; }}
            .toc-section {{ page-break-after: always; }}
            .section {{ page-break-inside: avoid; }}
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {content}
    <div class="export-info">
        <p>文档导出时间：{export_time}</p>
    </div>
</body>
</html>
"""
    
    def _load_export_templates(self) -> Dict[str, Dict[str, Any]]:
        """加载导出模板配置"""
        return {
            'html': {
                'template': self._get_default_html_template(),
                'description': '默认HTML导出模板'
            },
            'pdf': {
                'page_size': 'A4',
                'margin': '2cm',
                'font_family': 'Microsoft YaHei',
                'description': 'PDF导出配置'
            },
            'word': {
                'font_name': 'Microsoft YaHei',
                'font_size': 12,
                'line_spacing': 1.5,
                'description': 'Word导出配置'
            }
        }
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的导出格式列表"""
        return [fmt.value for fmt in self.supported_formats]
    
    def get_export_filename(self, project_name: str, export_format: ExportFormat, 
                          is_single_doc: bool = False, doc_title: str = "") -> str:
        """生成导出文件名"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if is_single_doc and doc_title:
            # 清理文件名中的特殊字符
            clean_title = "".join(c for c in doc_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"{clean_title}_{timestamp}.{export_format.value}"
        else:
            # 清理项目名中的特殊字符
            clean_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"{clean_name}_标书_{timestamp}.{export_format.value}"
        
        return filename