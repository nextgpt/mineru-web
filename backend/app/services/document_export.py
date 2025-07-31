"""
文档导出服务
支持PDF和Word格式的标书导出
"""
import io
import logging
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.models.tender import (
    TenderProject, 
    AnalysisResult, 
    OutlineStructure, 
    ChapterContent,
    DocumentInfo
)
from app.services.tender_storage import TenderStorageService
from app.services.document_templates import (
    template_manager, 
    document_formatter,
    StandardPDFTemplate,
    StandardWordTemplate
)
from app.services.document_management import DocumentManagementService, DocumentAccessLevel

logger = logging.getLogger(__name__)



class DocumentExportService:
    """文档导出服务"""
    
    def __init__(self, storage_service: TenderStorageService, document_manager: DocumentManagementService = None):
        self.storage = storage_service
        self.template_manager = template_manager
        self.formatter = document_formatter
        self.document_manager = document_manager
    
    async def export_to_pdf(
        self, 
        project: TenderProject, 
        template_name: str = "standard"
    ) -> Tuple[bytes, str]:
        """
        导出PDF格式标书
        
        Args:
            project: 项目对象
            template_name: 模板名称
            
        Returns:
            Tuple[bytes, str]: (PDF内容, 文件名)
        """
        try:
            # 加载项目数据
            analysis = await self.storage.load_analysis_result(project)
            outline = await self.storage.load_outline(project)
            chapters = await self.storage.load_all_chapters(project)
            
            # 创建PDF文档
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # 获取模板
            template = self.template_manager.get_pdf_template(template_name)
            
            # 构建文档内容
            story = []
            
            # 添加标题页
            story.extend(self._build_pdf_title_page(project, analysis, template))
            story.append(PageBreak())
            
            # 添加目录
            story.extend(self._build_pdf_table_of_contents(outline, template))
            story.append(PageBreak())
            
            # 添加章节内容
            for chapter in chapters:
                story.extend(self._build_pdf_chapter(chapter, template))
                story.append(Spacer(1, 12))
            
            # 生成PDF
            doc.build(story)
            pdf_content = buffer.getvalue()
            buffer.close()
            
            # 生成文件名
            filename = f"{project.project_name}_标书_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            logger.info(f"PDF导出完成，项目: {project.id}, 大小: {len(pdf_content)} bytes")
            return pdf_content, filename
            
        except Exception as e:
            logger.error(f"PDF导出失败，项目: {project.id}, 错误: {str(e)}")
            raise
    
    async def export_to_docx(
        self, 
        project: TenderProject, 
        template_name: str = "standard"
    ) -> Tuple[bytes, str]:
        """
        导出Word格式标书
        
        Args:
            project: 项目对象
            template_name: 模板名称
            
        Returns:
            Tuple[bytes, str]: (Word内容, 文件名)
        """
        try:
            # 加载项目数据
            analysis = await self.storage.load_analysis_result(project)
            outline = await self.storage.load_outline(project)
            chapters = await self.storage.load_all_chapters(project)
            
            # 获取模板并创建Word文档
            template = self.template_manager.get_word_template(template_name)
            doc = Document()
            template.setup_document_styles(doc)
            
            # 添加标题页
            self._build_word_title_page(doc, project, analysis)
            doc.add_page_break()
            
            # 添加目录
            self._build_word_table_of_contents(doc, outline)
            doc.add_page_break()
            
            # 添加章节内容
            for chapter in chapters:
                self._build_word_chapter(doc, chapter)
                doc.add_paragraph()  # 添加段落间距
            
            # 保存到内存
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_content = buffer.getvalue()
            buffer.close()
            
            # 生成文件名
            filename = f"{project.project_name}_标书_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            logger.info(f"Word导出完成，项目: {project.id}, 大小: {len(docx_content)} bytes")
            return docx_content, filename
            
        except Exception as e:
            logger.error(f"Word导出失败，项目: {project.id}, 错误: {str(e)}")
            raise
    
    async def generate_download_url(
        self, 
        project: TenderProject, 
        document_path: str, 
        expires: int = 3600
    ) -> str:
        """
        生成安全下载链接
        
        Args:
            project: 项目对象
            document_path: 文档路径
            expires: 过期时间（秒）
            
        Returns:
            str: 下载链接
        """
        try:
            # 验证文档路径是否属于该项目
            if not document_path.startswith(project.get_storage_path()):
                raise ValueError("文档路径不属于该项目")
            
            # 生成预签名URL
            download_url = await self.storage.get_document_download_url(document_path, expires)
            
            logger.info(f"生成下载链接，项目: {project.id}, 路径: {document_path}")
            return download_url
            
        except Exception as e:
            logger.error(f"生成下载链接失败，项目: {project.id}, 错误: {str(e)}")
            raise
    
    async def export_document(
        self, 
        project: TenderProject, 
        format_type: str, 
        template_name: str = "standard",
        user_id: str = None,
        description: str = ""
    ) -> DocumentInfo:
        """
        导出文档并保存到存储
        
        Args:
            project: 项目对象
            format_type: 文档格式 (pdf, docx)
            template_name: 模板名称
            user_id: 用户ID
            description: 版本描述
            
        Returns:
            DocumentInfo: 文档信息
        """
        try:
            # 根据格式类型导出文档
            if format_type.lower() == 'pdf':
                content, filename = await self.export_to_pdf(project, template_name)
            elif format_type.lower() == 'docx':
                content, filename = await self.export_to_docx(project, template_name)
            else:
                raise ValueError(f"不支持的文档格式: {format_type}")
            
            # 如果有文档管理器，使用版本管理功能
            if self.document_manager and user_id:
                metadata = await self.document_manager.save_document_with_version(
                    project=project,
                    document_content=content,
                    filename=filename,
                    document_type=format_type.lower(),
                    created_by=user_id,
                    description=description,
                    access_level=DocumentAccessLevel.TENANT
                )
                
                # 转换为DocumentInfo格式
                current_version = next((v for v in metadata.versions if v.is_current), None)
                document_info = DocumentInfo(
                    document_id=metadata.document_id,
                    document_type=metadata.document_type,
                    filename=metadata.filename,
                    file_size=current_version.file_size if current_version else len(content),
                    minio_path=current_version.document_path if current_version else "",
                    created_at=metadata.created_at
                )
            else:
                # 使用传统方式保存
                document_path = await self.storage.save_document(
                    project, format_type, content, filename
                )
                
                # 创建文档信息
                document_info = DocumentInfo(
                    document_id=str(uuid.uuid4()),
                    document_type=format_type.lower(),
                    filename=filename,
                    file_size=len(content),
                    minio_path=document_path,
                    created_at=datetime.utcnow()
                )
                
                # 保存文档信息到项目元数据
                await self._save_document_info(project, document_info)
            
            logger.info(f"文档导出并保存完成，项目: {project.id}, 格式: {format_type}")
            return document_info
            
        except Exception as e:
            logger.error(f"文档导出失败，项目: {project.id}, 格式: {format_type}, 错误: {str(e)}")
            raise
    
    def _build_pdf_title_page(
        self, 
        project: TenderProject, 
        analysis: AnalysisResult, 
        template
    ) -> List:
        """构建PDF标题页"""
        story = []
        
        # 项目标题
        title = analysis.project_info.get('project_name', project.project_name)
        story.append(Paragraph(title, template.title_style))
        story.append(Spacer(1, 30))
        
        # 副标题
        story.append(Paragraph("技术方案书", template.chapter_title_style))
        story.append(Spacer(1, 50))
        
        # 使用格式化器生成项目信息表格
        project_data = self.formatter.format_project_info_table(analysis)
        
        table = Table(project_data, colWidths=[2*inch, 4*inch])
        table.setStyle(template.get_table_style(header_bg=True))
        
        story.append(table)
        return story
    
    def _build_pdf_table_of_contents(
        self, 
        outline: OutlineStructure, 
        template
    ) -> List:
        """构建PDF目录"""
        story = []
        
        story.append(Paragraph("目录", template.title_style))
        story.append(Spacer(1, 20))
        
        for chapter in outline.chapters:
            # 章节标题
            toc_entry = f"{chapter.chapter_id}. {chapter.title}"
            story.append(Paragraph(toc_entry, template.toc_title_style))
            
            # 子章节
            for subsection in chapter.subsections:
                sub_entry = f"    {subsection.get('id', '')} {subsection.get('title', '')}"
                story.append(Paragraph(sub_entry, template.toc_item_style))
        
        return story
    
    def _build_pdf_chapter(
        self, 
        chapter: ChapterContent, 
        template
    ) -> List:
        """构建PDF章节"""
        story = []
        
        # 章节标题
        story.append(Paragraph(f"{chapter.chapter_id}. {chapter.chapter_title}", template.chapter_title_style))
        story.append(Spacer(1, 12))
        
        # 章节内容 - 使用格式化器分割段落
        paragraphs = self.formatter.split_content_into_paragraphs(chapter.content)
        for paragraph in paragraphs:
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), template.body_style))
                story.append(Spacer(1, 6))
        
        return story
    
    def _build_word_title_page(
        self, 
        doc: Document, 
        project: TenderProject, 
        analysis: AnalysisResult
    ):
        """构建Word标题页"""
        # 项目标题
        title = analysis.project_info.get('project_name', project.project_name)
        title_para = doc.add_paragraph(title, style='Document Title')
        
        # 副标题
        doc.add_paragraph('技术方案书', style='Chapter Title')
        
        # 添加空行
        doc.add_paragraph()
        
        # 使用格式化器生成项目信息表格
        table_data = self.formatter.format_project_info_table(analysis)
        
        # 创建表格
        table = doc.add_table(rows=len(table_data), cols=2)
        table.style = 'Table Grid'
        
        for i, (label, value) in enumerate(table_data):
            table.cell(i, 0).text = str(label)
            table.cell(i, 1).text = str(value)
    
    def _build_word_table_of_contents(self, doc: Document, outline: OutlineStructure):
        """构建Word目录"""
        doc.add_paragraph('目录', style='Document Title')
        doc.add_paragraph()
        
        for chapter in outline.chapters:
            # 章节标题
            toc_entry = f"{chapter.chapter_id}. {chapter.title}"
            doc.add_paragraph(toc_entry, style='TOC Title')
            
            # 子章节
            for subsection in chapter.subsections:
                sub_entry = f"    {subsection.get('id', '')} {subsection.get('title', '')}"
                doc.add_paragraph(sub_entry, style='TOC Item')
    
    def _build_word_chapter(self, doc: Document, chapter: ChapterContent):
        """构建Word章节"""
        # 章节标题
        doc.add_paragraph(f"{chapter.chapter_id}. {chapter.chapter_title}", style='Chapter Title')
        
        # 章节内容 - 使用格式化器分割段落
        paragraphs = self.formatter.split_content_into_paragraphs(chapter.content)
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip(), style='Body Text')
    
    async def _save_document_info(self, project: TenderProject, document_info: DocumentInfo):
        """保存文档信息到项目元数据"""
        try:
            # 加载现有文档列表
            documents_path = f"{project.get_storage_path()}/documents_info.json"
            
            try:
                documents_data = await self.storage._load_json(documents_path)
                documents = [DocumentInfo(**doc) for doc in documents_data.get('documents', [])]
            except FileNotFoundError:
                documents = []
            
            # 添加新文档信息
            documents.append(document_info)
            
            # 保存更新后的文档列表
            documents_data = {
                'documents': [doc.dict() for doc in documents],
                'updated_at': datetime.utcnow().isoformat()
            }
            
            await self.storage._save_json(documents_path, documents_data)
            
        except Exception as e:
            logger.error(f"保存文档信息失败: {str(e)}")
            # 不抛出异常，避免影响主要的导出流程
    
    async def list_project_documents(self, project: TenderProject) -> List[DocumentInfo]:
        """列出项目的所有文档"""
        try:
            # 从存储服务获取文档列表
            return await self.storage.list_project_documents(project)
        except Exception as e:
            logger.error(f"列出项目文档失败: {str(e)}")
            return []
    
    async def delete_document(self, project: TenderProject, document_path: str) -> bool:
        """删除文档"""
        try:
            # 验证文档路径是否属于该项目
            if not document_path.startswith(project.get_storage_path()):
                raise ValueError("文档路径不属于该项目")
            
            # 删除文档文件
            self.storage.minio_client.remove_object(self.storage.bucket_name, document_path)
            
            logger.info(f"文档删除成功: {document_path}")
            return True
            
        except Exception as e:
            logger.error(f"删除文档失败: {str(e)}")
            return False
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的导出格式"""
        return ['pdf', 'docx']
    
    def get_available_templates(self) -> Dict[str, List[str]]:
        """获取可用的模板"""
        return self.template_manager.list_available_templates()